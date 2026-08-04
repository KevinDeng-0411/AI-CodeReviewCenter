"""C5 检索质量门禁：DOCX/PDF 分块 before/after 结构对比。

自包含（不修改共享 golden_retrieval.py，不依赖 PG/embedding）。
对比旧链路（str(e.text) 平铺 / pypdf 纯文本）与新链路（元素序列化 / pdfminer 字号标题）
在 SemanticChunker 下的分块结构：新链路应按标题切出更多块且各节内容分离。

FakeEmbedder 是 hash-based（同文本 sim≈1，不同文本近正交，不做语义匹配），R@5 无意义；
故用结构对比（chunk 数 + 节隔离）作为 before/after 量化证据。
"""

import io

from docx import Document as DocxDocument
from unstructured.partition.auto import partition

from app.ai.rag.semantic_chunker import SemanticChunker
from app.ai.services.document_parser import DocumentParserService


def _old_serialize_docx(content: bytes) -> str:
    """旧链路：partition -> str(e.text) 拼接（丢弃 Title/ListItem 元素类型）。"""
    elements = partition(
        file=io.BytesIO(content),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    return "\n\n".join(
        str(e.text) for e in elements if getattr(e, "text", None) and e.text.strip()
    )


def _docx_two_sections() -> bytes:
    """两节 DOCX：标题1 + 正文1 + 标题2 + 正文2。"""
    doc = DocxDocument()
    doc.add_heading("缓存击穿")
    doc.add_paragraph("互斥锁与逻辑过期是缓存击穿的两种应对方案。" * 4)
    doc.add_heading("检索增强")
    doc.add_paragraph("混合检索与 RRF 融合是检索增强的核心机制。" * 4)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_c5_docx_chunking_before_after():
    """DOCX：新链路按标题切出更多块，且两节内容不在同一块（标题边界生效）。"""
    chunker = SemanticChunker()
    content = _docx_two_sections()

    old_text = _old_serialize_docx(content)
    new_text = DocumentParserService()._parse_sync(content, "x.docx")

    old_chunks = chunker.chunk(old_text, content_type="md")
    new_chunks = chunker.chunk(new_text, content_type="md")

    print(f"[C5 EVAL] DOCX old_chunks={len(old_chunks)} new_chunks={len(new_chunks)}")
    print(f"[C5 EVAL] OLD  前 60 字: {old_text[:60]!r}")
    print(f"[C5 EVAL] NEW  前 60 字: {new_text[:60]!r}")

    # 1) 新链路按标题切出更多块（旧链路无 `#` -> 整篇一坨）
    assert len(new_chunks) > len(old_chunks), (
        f"new {len(new_chunks)} !> old {len(old_chunks)}"
    )

    # 2) 新链路：两节内容分离，不在同一块（标题边界真正生效）
    for c in new_chunks:
        assert not ("互斥锁" in c and "RRF" in c), (
            "两节内容混在同一块，标题边界未生效"
        )

    # 3) 两节各自出现在某块中
    assert any("缓存击穿" in c or "互斥锁" in c for c in new_chunks)
    assert any("检索增强" in c or "RRF" in c for c in new_chunks)
    print("[C5 EVAL] DOCX PASS: 标题边界生效，两节内容分离到不同块")


def test_c5_docx_new_chunks_carry_title_prefix():
    """新链路序列化后，标题在 chunk 文本中以 `# ` 开头（chunk_by_title 依赖的边界）。"""
    content = _docx_two_sections()
    new_text = DocumentParserService()._parse_sync(content, "x.docx")
    assert "# 缓存击穿" in new_text
    assert "# 检索增强" in new_text
