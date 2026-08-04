"""C1-C：文件类型策略与轻量文档解析。"""

import io

import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter

from app.ai.services.document_parser import DocumentParserService


def test_upload_type_policy_matches_supported_extensions_and_mime():
    parser = DocumentParserService()
    assert parser.supports_upload("readme.md", "text/markdown")
    assert parser.supports_upload("README.MD", "text/plain; charset=utf-8")
    assert parser.supports_upload("manual.pdf", "application/pdf")
    assert parser.supports_upload("manual.docx", None)
    assert parser.supports_upload("manual.docx", "application/octet-stream")
    assert not parser.supports_upload("manual.pdf", "image/png")
    assert not parser.supports_upload("archive.zip", "application/octet-stream")
    assert not parser.supports_upload("no-extension", "text/plain")


def test_safe_filename_removes_client_paths():
    parser = DocumentParserService()
    assert parser.safe_filename("../../secret/readme.md") == "readme.md"
    assert parser.safe_filename(r"C:\fakepath\manual.docx") == "manual.docx"
    assert parser.safe_filename(None) == ""


def test_unknown_extension_does_not_fall_back_to_plain_text():
    with pytest.raises(ValueError, match="unsupported document type"):
        DocumentParserService._guess_content_type("payload.exe")


@pytest.mark.parametrize(
    ("filename", "content", "expected"),
    [
        ("notes.txt", b"first line\nsecond line", "second line"),
        ("README.md", b"# Heading\n\nMarkdown body", "Markdown body"),
        ("page.html", b"<html><body><h1>Title</h1><p>HTML body</p></body></html>", "HTML body"),
    ],
)
async def test_parse_small_text_formats(filename, content, expected):
    text = await DocumentParserService().parse(content, filename)
    assert expected in text


async def test_parse_docx_with_declared_extra():
    document = DocxDocument()
    document.add_heading("DOCX heading")
    document.add_paragraph("DOCX body")
    buffer = io.BytesIO()
    document.save(buffer)

    text = await DocumentParserService().parse(buffer.getvalue(), "manual.docx")
    assert "DOCX heading" in text
    assert "DOCX body" in text


async def test_pdf_without_text_layer_returns_empty_for_route_to_reject():
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    buffer = io.BytesIO()
    writer.write(buffer)

    text = await DocumentParserService().parse(buffer.getvalue(), "scan.pdf")
    assert text == ""


# ---------- C5：元素感知序列化 + PDF 字号标题 + parse->chunker 集成 ----------


def _pdf_with_title_and_body(title: str, body: str) -> bytes:
    """pypdf 生成带标题(18pt)+正文(12pt)两个文本流的合法 PDF。"""
    from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): writer._add_object(font)})}
    )
    stream = DecodedStreamObject()
    stream.set_data(
        f"BT /F1 18 Tf 72 720 Td ({title}) Tj ET"
        f" BT /F1 12 Tf 72 700 Td ({body}) Tj ET".encode()
    )
    page[NameObject("/Contents")] = writer._add_object(stream)
    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


async def test_c5_markdown_title_serialized_back_to_hash():
    """md 的 Title 元素经 _serialize 补回 `# `，chunker 二次解析仍识别为标题。"""
    text = await DocumentParserService().parse("# 缓存方案\n\n正文内容".encode(), "x.md")
    assert text.startswith("# 缓存方案")


async def test_c5_docx_heading_and_list_serialized():
    """DOCX heading -> `# `、List Bullet -> `- `（元素类型穿到序列化层）。"""
    document = DocxDocument()
    document.add_heading("DOCX 标题")
    document.add_paragraph("正文段落")
    document.add_paragraph("列表项", style="List Bullet")
    buffer = io.BytesIO()
    document.save(buffer)

    text = await DocumentParserService().parse(buffer.getvalue(), "manual.docx")
    assert "# DOCX 标题" in text
    assert "正文段落" in text
    assert "- 列表项" in text


async def test_c5_pdf_text_layer_title_detected_by_font_size():
    """文本层 PDF：标题(18pt)字号 > 正文(12pt) -> 检测为 `# ` 标题。"""
    pdf = _pdf_with_title_and_body("Section One", "This is the body paragraph about caching.")
    text = await DocumentParserService().parse(pdf, "doc.pdf")
    assert "# Section One" in text
    assert "body paragraph" in text


async def test_c5_parse_then_chunk_splits_on_titles():
    """端到端：parse(DOCX 多标题) -> SemanticChunker 按 `#` 边界切出 >=2 块。"""
    from app.ai.rag.semantic_chunker import SemanticChunker

    document = DocxDocument()
    document.add_heading("第一章 缓存")
    document.add_paragraph("缓存击穿与雪崩方案。" * 5)
    document.add_heading("第二章 检索")
    document.add_paragraph("混合检索与 RRF 融合。" * 5)
    buffer = io.BytesIO()
    document.save(buffer)

    text = await DocumentParserService().parse(buffer.getvalue(), "manual.docx")
    chunks = SemanticChunker().chunk(text, content_type="md")
    assert len(chunks) >= 2, f"expected >=2 chunks by title, got {len(chunks)}: {chunks}"
    # 两个标题都应作为某块的开头出现
    joined = "\n".join(chunks)
    assert "第一章 缓存" in joined
    assert "第二章 检索" in joined
